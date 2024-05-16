import docx
import streamlit as st
import json
from io import BytesIO
import openai
from openai import AzureOpenAI

client = AzureOpenAI(
   api_key = '1767ca3c571d462f93178c82c4308a9f',
   api_version = "2023-05-15",
   azure_endpoint = "https://popup.openai.azure.com/"
)

# Setting page title and header
st.set_page_config(page_title="Azure OpenAI ChatGPT", page_icon=":robot_face:")
st.markdown("<h1 style='text-align: center;'>Test case generator</h1>", unsafe_allow_html=True)

def extract_text_from_docx(file_content):
  doc = docx.Document(BytesIO(file_content))
  extracted_data = {}
  current_heading = None
  table_headings = ["Data Input Details:","Data Output Details:"]

  for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Normal'):
      current_heading = paragraph.text
      extracted_data[current_heading] = []
    else:
      extracted_data[current_heading].append(paragraph.text)

  for index, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
      row_data = []
      for cell in row.cells:
        row_data.append(cell.text)
      table_data.append(row_data)
    extracted_data[table_headings[index]] = table_data
  return extracted_data

@st.cache_data
def create_context(file_content):
    data = extract_text_from_docx(file_content)
    context = json.dumps(data,sort_keys=True)
    return context

@st.cache_data
def callAIforCases(selected_options, context):
    mapper = {'Positive':'positive','Negative':'negative', 'Edge-Possibilities':'edge cases'}
    selected_string = mapper[selected_options[0]]
    for element in selected_options[1:]:
        selected_string = selected_string + ", " + mapper[element]

    prompt = "Generate test cases scenarios for the given Context divided into the groups of " + selected_string + ". No need of test input details and test output details. The output should be in form of having heading of scenarios without any extra text"

    response = client.chat.completions.create(
                                            model = 'project',
                                            messages=[
                                                {"role":"system", "content":"You are a helpful assistant"},
                                                {"role":"user", "content":"Context: " + context +" Query: " + prompt}
                                            ]
                                        )
    return response.choices[0].message.content.split('\n')

@st.cache_data
def callAIfortestdata(selected_case, count_of_cases, context):
    
    prompt = "Generate " + count_of_cases + " examples of test data for " + selected_case + "scenario. Generate the data in above format given."

    response = client.chat.completions.create(
                                            model = 'project',
                                            messages=[
                                                {"role":"system", "content":"You are a helpful assistant"},
                                                {"role":"user", "content":"Context: " + context +" Query: " + prompt}
                                            ]
                                        )
    return response.choices[0].message.content.split('\n')

# Disable generate cases button if all field are not filled
@st.cache_data
def update_button_state(uploaded_file, selected_options, count_of_cases):
    return uploaded_file is not None and selected_options and count_of_cases

def main():
    uploaded_file = st.sidebar.file_uploader("Choose your file:", type=[".docx"])
    context = ""
    if uploaded_file is not None:
        context = create_context(uploaded_file.read())
        st.sidebar.success("File uploaded successfully!")
    else:
        st.sidebar.write("File not uploaded")

    headings=['Positive','Negative','Edge-Possibilities']
    selected_options=st.sidebar.multiselect('Select type of cases: ',headings, default=headings)

    count_of_cases = str(st.sidebar.number_input("Choose number of test data to be generated",min_value=1, max_value=10, step=1))
    
    if 'is_generate_clicked' not in st.session_state:
        st.session_state.is_generate_clicked = None

    all_fields_filled = update_button_state(uploaded_file, selected_options, count_of_cases)
    temp = st.button('Generate Cases', disabled = not all_fields_filled)
    
    if st.session_state.is_generate_clicked is None and temp:
        st.session_state.is_generate_clicked = temp

    if not all_fields_filled:
        st.write("Please fill all fields to enable the Generate button.")
    else:
        if st.session_state.is_generate_clicked:
            test_cases = callAIforCases(selected_options, context)
            heading = True
            for i, test_case in enumerate(test_cases):
                button_key=f"button_{i}"
                if test_case == "":
                    heading = True
                    continue
                if heading:
                    st.markdown(f"<h3>{test_case}</h3>",unsafe_allow_html=True)
                    heading = False
                    continue
                if st.button(test_case, key = button_key):
                    selected_case = test_case
                    #data_slot = st.empty()
                    data = callAIfortestdata(selected_case,count_of_cases, context)
                    st.write(data)

if __name__ == "__main__":
    main()